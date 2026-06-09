"""Per-engagement client-document ingestion and search for Phase agents.

Documents are chunked locally, embedded with OpenAI ``text-embedding-3-large``
(3072 dimensions), and stored in a dense-only Qdrant collection named
``client_docs_{engagement_id}``. Empty/offline calls are safe no-ops so tests and
demo runs do not need Qdrant or OpenAI credentials.
"""
from __future__ import annotations

import hashlib
import io
import json
import logging
import os
import re
import tempfile
import zipfile
import base64
from pathlib import PurePosixPath
from typing import Any
from xml.etree import ElementTree

from aaa.platform.evidence import EvidenceStore
from aaa.settings import settings

logger = logging.getLogger(__name__)

_DENSE_MODEL = "text-embedding-3-large"
_DENSE_DIM = 3072
_VECTOR_NAME = "dense"
_CHUNK_CHARS = 1600
_OVERLAP_CHARS = 200
_EMBED_BATCH = 64
_UPSERT_BATCH = 128
_OFFLINE = os.environ.get("AAA_OFFLINE_MODE", "false").lower() == "true"


def _openai_api_key() -> str:
    return os.environ.get("OPENAI_API_KEY") or settings.openai_api_key


def _embeddings_available() -> bool:
    return bool(_openai_api_key())


def _collection_name(engagement_id: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9_]", "_", engagement_id.replace("-", "_"))
    return f"client_docs_{safe}"


def _filename(uri: str) -> str:
    return PurePosixPath(uri.split("?", 1)[0]).name or "client_document"


def _content_type(uri: str) -> str:
    suffix = PurePosixPath(uri.split("?", 1)[0]).suffix.lower().lstrip(".")
    return suffix if suffix in {"pdf", "docx", "txt", "json", "md"} else "txt"


def _document_role(uri: str) -> str:
    name = _filename(uri).lower()
    role_markers = {
        "risk_management_file": ("risk_management", "risk-management"),
        "post_market_plan": ("post_market", "post-market"),
        "eu_declaration": ("eu_doc", "eu_declaration", "declaration"),
        "system_prompt": ("system_prompt", "system-prompt"),
        "rag_manifest": ("rag_manifest", "rag-manifest"),
        "guardrail_config": ("guardrail",),
        "golden_set": ("golden_set", "golden-set"),
    }
    for role, markers in role_markers.items():
        if any(marker in name for marker in markers):
            return role
    return "unknown"


def _coerce_bytes(content: Any) -> bytes:
    if isinstance(content, dict) and isinstance(content.get("body_base64"), str):
        return base64.b64decode(content["body_base64"])
    if isinstance(content, bytes):
        return content
    if isinstance(content, bytearray):
        return bytes(content)
    if isinstance(content, str):
        return content.encode("utf-8")
    return json.dumps(content, indent=2, default=str).encode("utf-8")


def _load_document(uri: str, store: EvidenceStore | None) -> bytes | None:
    if uri.startswith("minio://") and store is not None:
        content = store.get_artefact(uri)
        return None if content is None else _coerce_bytes(content)
    return None


def _extract_pdf_pages(data: bytes) -> list[tuple[int | None, str]]:
    import pypdfium2  # type: ignore

    pages: list[tuple[int | None, str]] = []
    with tempfile.NamedTemporaryFile(suffix=".pdf") as tmp:
        tmp.write(data)
        tmp.flush()
        doc = pypdfium2.PdfDocument(tmp.name)
        try:
            for idx, page in enumerate(doc, start=1):
                textpage = page.get_textpage()
                try:
                    pages.append((idx, textpage.get_text_bounded() or ""))
                finally:
                    textpage.close()
                    page.close()
        finally:
            doc.close()
    return pages


def _extract_docx_pages(data: bytes) -> list[tuple[int | None, str]]:
    try:
        import docx  # type: ignore

        document = docx.Document(io.BytesIO(data))
        text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
        return [(None, text)]
    except ImportError:
        pass

    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        xml = archive.read("word/document.xml")
    root = ElementTree.fromstring(xml)
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    text = "\n".join(node.text or "" for node in root.iter(f"{namespace}t"))
    return [(None, text)]


def _extract_pages(data: bytes, content_type: str) -> list[tuple[int | None, str]]:
    if content_type == "pdf":
        return _extract_pdf_pages(data)
    if content_type == "docx":
        return _extract_docx_pages(data)
    return [(None, data.decode("utf-8", errors="replace"))]


def _section_hint(text: str) -> str | None:
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("#") or stripped[:1].isdigit():
            return stripped[:120]
    return None


def _chunks_for_document(uri: str, data: bytes) -> list[dict[str, Any]]:
    ctype = _content_type(uri)
    source_sha = hashlib.sha256(data).hexdigest()
    chunks: list[dict[str, Any]] = []
    step = max(1, _CHUNK_CHARS - _OVERLAP_CHARS)
    for page_number, text in _extract_pages(data, ctype):
        text = re.sub(r"\s+", " ", text).strip()
        if not text:
            continue
        for start in range(0, len(text), step):
            body = text[start:start + _CHUNK_CHARS].strip()
            if not body:
                continue
            chunks.append({
                "text": body,
                "source_uri": uri,
                "source_filename": _filename(uri),
                "source_sha256": source_sha,
                "content_type": ctype,
                "document_role": _document_role(uri),
                "page_number": page_number,
                "section_hint": _section_hint(body),
                "char_start": start,
                "char_end": min(start + _CHUNK_CHARS, len(text)),
            })
    total = len(chunks)
    for idx, chunk in enumerate(chunks):
        chunk["chunk_index"] = idx
        chunk["chunk_total"] = total
    return chunks


def _qdrant_client() -> Any:
    import qdrant_client

    return qdrant_client.QdrantClient(
        url=os.environ.get("QDRANT_URL", "http://localhost:6333"),
        api_key=os.environ.get("QDRANT_API_KEY") or None,
        prefer_grpc=False,
        check_compatibility=False,
    )


def _ensure_collection(client: Any, collection: str) -> None:
    from qdrant_client import models as qmodels

    if not client.collection_exists(collection):
        client.create_collection(
            collection_name=collection,
            vectors_config={
                _VECTOR_NAME: qmodels.VectorParams(
                    size=_DENSE_DIM, distance=qmodels.Distance.COSINE,
                )
            },
        )
    for key in ("source_uri", "document_role", "source_sha256"):
        try:
            client.create_payload_index(
                collection_name=collection,
                field_name=key,
                field_schema=qmodels.PayloadSchemaType.KEYWORD,
            )
        except Exception:  # pragma: no cover - index already exists
            pass


def _source_exists(client: Any, collection: str, source_uri: str) -> bool:
    from qdrant_client import models as qmodels

    points, _ = client.scroll(
        collection_name=collection,
        scroll_filter=qmodels.Filter(
            must=[qmodels.FieldCondition(
                key="source_uri", match=qmodels.MatchValue(value=source_uri),
            )]
        ),
        limit=1,
        with_payload=False,
        with_vectors=False,
    )
    return bool(points)


def _embed(texts: list[str]) -> list[list[float]]:
    import openai

    client = openai.OpenAI(api_key=_openai_api_key())
    vectors: list[list[float]] = []
    for start in range(0, len(texts), _EMBED_BATCH):
        batch = texts[start:start + _EMBED_BATCH]
        resp = client.embeddings.create(model=_DENSE_MODEL, input=batch)
        vectors.extend([list(item.embedding) for item in resp.data])
    return vectors


def _point_id(chunk: dict[str, Any]) -> str:
    raw = f"{chunk['source_uri']}:{chunk['chunk_index']}:{chunk['source_sha256']}"
    digest = hashlib.sha256(raw.encode("utf-8")).hexdigest()
    return f"{digest[:8]}-{digest[8:12]}-{digest[12:16]}-{digest[16:20]}-{digest[20:32]}"


def client_doc_ingest(
    engagement_id: str,
    doc_uris: list[str],
    store: EvidenceStore | None = None,
) -> dict[str, Any]:
    """Ingest client-uploaded documents into an engagement-specific Qdrant collection."""
    collection = _collection_name(engagement_id)
    if not doc_uris:
        return {"collection_name": collection, "chunks_indexed": 0, "sources": []}
    if _OFFLINE or not _embeddings_available():
        if doc_uris and not _OFFLINE:
            logger.info(
                "Skipping client_doc_ingest for %s because OPENAI_API_KEY is not configured.",
                engagement_id,
            )
        return {"collection_name": collection, "chunks_indexed": 0, "sources": []}

    client = _qdrant_client()
    _ensure_collection(client, collection)
    chunks: list[dict[str, Any]] = []
    sources: list[str] = []
    for uri in doc_uris:
        if _source_exists(client, collection, uri):
            sources.append(uri)
            continue
        data = _load_document(uri, store)
        if data is None:
            logger.warning("Client document not found for ingestion: %s", uri)
            continue
        doc_chunks = _chunks_for_document(uri, data)
        chunks.extend(doc_chunks)
        if doc_chunks:
            sources.append(uri)

    if not chunks:
        return {"collection_name": collection, "chunks_indexed": 0, "sources": sources}

    from qdrant_client import models as qmodels

    vectors = _embed([chunk["text"] for chunk in chunks])
    written = 0
    for start in range(0, len(chunks), _UPSERT_BATCH):
        batch = chunks[start:start + _UPSERT_BATCH]
        batch_vectors = vectors[start:start + _UPSERT_BATCH]
        points = [
            qmodels.PointStruct(
                id=_point_id(chunk), payload=chunk, vector={_VECTOR_NAME: vector},
            )
            for chunk, vector in zip(batch, batch_vectors, strict=True)
        ]
        client.upsert(collection_name=collection, points=points, wait=True)
        written += len(points)
    return {"collection_name": collection, "chunks_indexed": written, "sources": sources}


def _collection_exists(client: Any, collection: str) -> bool:
    try:
        return bool(client.collection_exists(collection))
    except Exception:
        return False


def _hit(point: Any) -> dict[str, Any]:
    payload = getattr(point, "payload", None) or {}
    return {
        "text": payload.get("text", ""),
        "source_uri": payload.get("source_uri", ""),
        "source_sha256": payload.get("source_sha256", ""),
        "document_role": payload.get("document_role", "unknown"),
        "page_number": payload.get("page_number"),
        "section_hint": payload.get("section_hint"),
        "chunk_index": payload.get("chunk_index", 0),
        "chunk_total": payload.get("chunk_total", 0),
        "score": float(getattr(point, "score", 0.0) or 0.0),
    }


def client_doc_search(engagement_id: str, query: str, top_k: int = 3) -> list[dict[str, Any]]:
    """Search an engagement's client-document collection; return [] if unavailable."""
    if _OFFLINE or not _embeddings_available():
        return []
    collection = _collection_name(engagement_id)
    try:
        client = _qdrant_client()
        if not _collection_exists(client, collection):
            return []
        vector = _embed([query])[0]
        response = client.query_points(
            collection_name=collection,
            query=vector,
            using=_VECTOR_NAME,
            limit=top_k,
            with_payload=True,
        )
        return [_hit(point) for point in response.points]
    except Exception as exc:  # pragma: no cover - offline/no-service fallback
        logger.warning("client_doc_search unavailable for %s: %s", engagement_id, exc)
        return []


__all__ = ["client_doc_ingest", "client_doc_search"]
